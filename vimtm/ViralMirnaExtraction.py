import pandas as pd
from docx import Document
import regex as re
import glob
from lxml import etree
import logging

from collections import Counter, defaultdict
import pdfplumber
import zipfile
from natsort import natsorted
from collections import OrderedDict

import os, sys
sys.path.insert(0, str(os.path.dirname(os.path.realpath(__file__))) + "/../miRExplore/python/")

from synonymes.GeneOntology import GeneOntology, GOTerm, GORelation, GORelationType

import tempfile
import os
import copy






#
##
### Global Definitions
##
#


fastaSequenceMirRE = re.compile(r'^[ATCGU]{15,30}$')
mirnaTextSearch = re.compile(r'^[ATCGU]{15,30}$')

mirnaFasta = re.compile(r'>.*\n[ATCGU]{15,30}$')
mirnaFastaMatch = re.compile(r'>(.*)\n([ATCGU]{15,30})\n')
fastaStartRE = re.compile(r'^>\S\S\S\S+')


#
##
### Classes
##
#

class DocumentMirnaSequence:

    def __init__(self, seqid, seq, source=None):

        self.mirna_id = seqid.replace("  ", " ").replace(" ", "_").replace(".", "_")
        self.mirna_seq = seq.replace("U", "T")
        self.origin = []
        self.source = source
        self.bestmatches  = {}

    def add_origin(self, neworigin):
        if not neworigin in self.origin:
            self.origin.append(str(neworigin))

    def to_fasta(self):
        return ">{}\n{}".format(self.mirna_id, self.mirna_seq.replace("U", "T"))

    def __str__(self):
        return "{} {} {} {}".format(self.mirna_id, self.mirna_seq, self.source, ",".join(self.origin))

    def __repr__(self):
        return "<{}/>".format(self.__str__())

    def __eq__(self, other):
        return isinstance(other, DocumentMirnaSequence) and self.source==other.source and self.mirna_id == other.mirna_id and self.mirna_seq == other.mirna_seq and tuple(self.origin) == tuple(other.origin)
    
    def __ne__(self, other):
        return not self.__eq__(other)

    def __hash__(self):
        return hash(self.mirna_id) + hash(self.mirna_seq) + hash(tuple(self.origin)) + hash(self.source)

    @property
    def matched_taxids(self):
        if self.bestmatches is None:
            return None

        result = defaultdict(set)

        for species in self.bestmatches:
            for hit in self.bestmatches[species]:
                result[species].add( hit["matched_tax_id"] )

        return result

    @property
    def matched_docids(self):
        if self.bestmatches is None:
            return None

        result = defaultdict(lambda: defaultdict(set))

        for species in self.bestmatches:
            for hit in self.bestmatches[species]:
                for hitTaxID in hit["matched_tax_id"]:
                    for dvirus in hit["matched_doc_virus"]:
                        result[species][hitTaxID].add( dvirus )

        return result



#
##
### MIRNA EXTRACTION
##
#

def is_empty(elem):
    if elem is None:
        return True

    if len(str(elem)) == 0:
        return True

    return False

def test_substrings_included(word, substrs):
    for x in substrs:
        if x in word:
            return True
    return False

def get_sequence(word):
    mirSequence = str(word).replace("Sequence: ", "").strip()
    mirSequence = re.sub(r"^[\s0-9]+","",mirSequence)

    return mirSequence

def test_is_sequence( word):
    mirSequence = get_sequence(word)
    match = fastaSequenceMirRE.search(mirSequence)

    if match:
        return True

    match = mirnaFasta.search(mirSequence)
    if match:
        return True

    return False
    
def hasSeqColumn(df, num_tests=10):
    hasSeqColumns = False
    for i in range(df.shape[1]):
        isSeqCol = column_contains_mirnas(df, 0, i, num_tests=num_tests)
        #print("col", i, isSeqCol)
        hasSeqColumns = hasSeqColumns or isSeqCol

    return hasSeqColumns

def column_contains_mirnacells(df, indexRow, xi, num_tests=10):
    

    foundSeqs = 0
    for ri, row in df.loc[indexRow+1:min(num_tests, df.shape[0]),].iterrows():
        
        cellContent = str(row[xi])

        #print('"'+cellContent+'"')
        #print(mirnaFasta.search(cellContent))

        if mirnaFasta.search(cellContent):
            foundSeqs += 1

    if foundSeqs > 0:
        return True

    #print("[column_contains_mirnacells] Column", xi, "is not a valid miRNA Sequence", foundSeqs)
    return False

def column_contains_mirnas(df, indexRow, xi, num_tests=10):

    foundSeqs = 0
    testedSeqs = 0
    for ri, row in df.loc[indexRow+1:min(num_tests, df.shape[0]),].iterrows():
        
        word = str(row[xi])

        #print(word, test_is_sequence(word))
        if not is_empty(word):
            testedSeqs+= 1

        if test_is_sequence(word):
            foundSeqs += 1
        #else:
        #    print("F", word)

    if testedSeqs>0 and foundSeqs/testedSeqs > 0.8:
        return True

    #print("[column_contains_mirnas] Column", xi, "is not a valid miRNA Sequence", foundSeqs, testedSeqs)
    return False

def column_contains_highlighted_mirnas(table, indexRow, xi):

    foundSeqs = 0
    testedSeqs = 0
    for ri, row in enumerate(table.rows):
        
        if ri <= indexRow:
            continue


        for par in row.cells[xi].paragraphs:
            cellWasFilled = False
            mirnaWasFound = False

            for run in par.runs:

                specialText = bool(run.underline) or bool(run.bold) or bool(run.italic)

                if not is_empty(run.text):
                    cellWasFilled = True
                
                if specialText and test_is_sequence(run.text):
                    mirnaWasFound = True
                    print("mirna", run.text)

            if cellWasFilled:
                testedSeqs += 1

            if mirnaWasFound: 
                foundSeqs += 1
        
    if  testedSeqs>0 and foundSeqs/testedSeqs > 0.8:
        return True

    print("[column_contains_highlighted_mirnas] Column", xi, "is not a valid miRNA Sequence", foundSeqs)
    return False

def getNameSeqColumns(df, table=None, num_tests=20):

    origDF = df.copy()

    for rowstart in range(0,5):

        df = origDF.copy()
        df = df.iloc[rowstart:]

        if 0 in df.shape:
            continue

        indexRow = 0
        for i in range(0, min(10, df.shape[0])):
            idxCnt = Counter(["str" if not pd.isna(x) and not x == "nan" else "NAN" for x in df.iloc[i]])
            #print(i,idxCnt)
            if idxCnt.most_common(1)[0][0] == "str":
                indexRow = i
                break

        #print(indexRow)
        potentialColumns = list(df.iloc[indexRow,])
        potentialColumns = [str(x).strip() for x in potentialColumns]
        #print("S", rowstart, potentialColumns)
        
        mirName = None
        mirSeq = None
        mirReason = None

        mirFounds = []

        for xi, x in enumerate(potentialColumns):
            if test_substrings_included(x.upper(), ["NAME","ACCESSION", "MATURE MIRNA", "MIRNA"]) and not test_substrings_included(x.upper(), ["TARGET GENE", "PROTEIN", "GENE"]) and not column_contains_mirnas(df, indexRow, xi, num_tests=num_tests):
                mirName = xi
                print("set mirName", mirName, mirSeq, potentialColumns[xi])

            elif test_substrings_included(x.upper(), ["MIRNA SEQUENCE","MATURE SEQUENCE", "SEQUENCE OF MATURE", "SEQUENCE", "5P MIRNA", "3P MIRNA"]) and not test_substrings_included(x.upper(), ["PRIMER"]) and column_contains_mirnas(df, indexRow, xi, num_tests=num_tests) and mirSeq is None:

                mirSeq = xi
                mirReason = "SEQ_COLUMN"
                print("set mirSeq", mirName, mirSeq, potentialColumns[xi])
                
            elif test_substrings_included(x.upper(), ["MIRNA"]) and column_contains_mirnacells(df, indexRow, xi, num_tests=num_tests) and mirSeq is None:
                mirSeq = xi
                mirName = xi
                mirReason = "FASTA_COLUMN"

            elif not table is None and test_substrings_included(x.upper(), ["PRECURSOR SEQUENCE", "MIRNA SEQUENCE","MATURE SEQUENCE", "SEQUENCE OF MATURE", "SEQUENCE"]) and column_contains_highlighted_mirnas(table, indexRow, xi) and mirSeq is None:
                mirSeq = xi
                mirReason = "HIGHLIGHT_COLUMN"

            if mirSeq != None and mirName != None:

                mirFounds.append((mirName, mirSeq, mirReason))

                #mirName = None
                mirSeq = None
                mirReason = None


        if mirName is None and mirSeq is None:
            #print(potentialColumns)

            for xi, x in enumerate(potentialColumns):

                if test_substrings_included(x.upper(), ["MIRNA", "MIRNA SEQUENCE","MATURE SEQUENCE", "SEQUENCE OF MATURE", "SEQUENCE"]) and column_contains_mirnas(df, indexRow, xi, num_tests=num_tests) and mirSeq is None:

                    mirSeq = xi
                    mirName = xi
                    mirReason = "NAME_SEQ_COLUMN"

                if mirSeq != None and mirName != None:

                    mirFounds.append((mirName, mirSeq, mirReason))
                    mirName = None
                    mirSeq = None
                    mirReason = None



        if len(mirFounds) > 0:
            for mirName, mirSeq, mirReason in mirFounds:
                print("Name:", mirName, potentialColumns[mirName] if mirName != None else "")
                print("Seq: ", mirSeq, potentialColumns[mirSeq] if mirSeq != None else "")
                print("Reason: ", mirReason)

            return indexRow, mirFounds

    print(mirFounds)
    for mirName, mirSeq, mirReason in mirFounds:
        print("Name:", mirName, potentialColumns[mirName] if mirName != None else "")
        print("Seq: ", mirSeq, potentialColumns[mirSeq] if mirSeq != None else "")
        print("Reason: ", mirReason)

    return indexRow, mirFounds

def extract_mirnas_from_document(infile):

    identified_mirnas = set()

    document = Document(infile)
    docTables = {}
    for table in document.tables:
        data = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(data)
        docTables[len(docTables)] = df
        
        identified_mirnas = identified_mirnas.union(extract_mirnas_from_pandas_tabledf( df, table ))

    return docTables,identified_mirnas

def extract_mirnas_from_pandas_tabledf( df, table=None, num_tests=10):
    startRow, mirFounds = getNameSeqColumns(df, table, num_tests)

    identified_mirnas = set()
    
    for mirName, mirSeq, mirReason in mirFounds:

        print(mirSeq, mirName, mirReason)
        if mirName == None or mirSeq == None:
            print("Skipping")
            print(startRow, mirName, mirSeq)
            continue



        if mirReason == "SEQ_COLUMN":

            for ri, row in df.loc[startRow:,].iterrows():

                mirSeqText = str(row[mirSeq])
                mirNameText = str(row[mirName])

                if not is_empty(mirSeqText) and not is_empty(mirSeqText):
                    mirSequence = get_sequence(mirSeqText)

                    if test_is_sequence(mirSeqText):
                        #print(mirNameText, mirSequence)
                        identified_mirnas.add( DocumentMirnaSequence(mirNameText, mirSequence, source="SEQ_COLUMN") )

        elif mirReason == "FASTA_COLUMN":

            for ri, row in df.loc[startRow:,].iterrows():
                cellContent = str(row[mirSeq])
                if mirnaFasta.search(cellContent):
                    for mirNameText, mirSequence in mirnaFastaMatch.findall(cellContent):
                        identified_mirnas.add( DocumentMirnaSequence(mirNameText, mirSequence, source="FASTA_COLUMN") )

        elif mirReason == "HIGHLIGHT_COLUMN":

            for ri, row in enumerate(table.rows):
        
                if ri <= startRow:
                    continue

                mirNameText = str(row.cells[mirName].text)

                cellMirnaCount = 0
                for par in row.cells[mirSeq].paragraphs:
                    cellWasFilled = False
                    mirnaWasFound = False

                    for run in par.runs:

                        specialText = bool(run.underline) or bool(run.bold) or bool(run.italic)

                        if not is_empty(run.text):
                            cellWasFilled = True
                        
                        runText = get_sequence(run.text)

                        if specialText and test_is_sequence(run.text):
                            identified_mirnas.add( DocumentMirnaSequence("{}_{}".format(mirNameText, cellMirnaCount), runText, source="HIGHLIGHT_COLUMN") )
                            cellMirnaCount += 1

        elif mirReason == "NAME_SEQ_COLUMN":

            mirSeqSequence = 0
            for ri, row in df.loc[startRow:,].iterrows():

                mirSeqText = str(row[mirSeq])

                if not is_empty(mirSeqText) and not is_empty(mirSeqText):
                    mirSequence = get_sequence(mirSeqText)

                    if test_is_sequence(mirSeqText):
                        #print(mirNameText, mirSequence)
                        identified_mirnas.add( DocumentMirnaSequence("miRNA_{}".format(mirSeqSequence), mirSequence, source="NAME_SEQ_COLUMN") )
                        mirSeqSequence += 1

        else:
            raise ValueError("Unknown mirReason")
       
    return identified_mirnas


def extract_mirnas_from_excel(infile):
    try:
        sheetsDF = pd.read_excel(infile, header=None, sheet_name=None, engine='openpyxl')
        print("Sheets", [x for x in sheetsDF])
    except:
        logging.error("Error loading file EXCEL: {}".format(infile))

        return None, None
    #columns is empty!

    identified_mirnas = extract_mirnas_sheetsdf(sheetsDF)
    return sheetsDF, identified_mirnas

def extract_mirnas_from_pdf(infile):
    pdfDFs = {}

    try:
        pdf = pdfplumber.open(infile)
    except:
        logging.error("Error loading file PDF: {}".format(infile))
        return None, None

    extractionSettings = OrderedDict()
    extractionSettings["loose"] = {
                "vertical_strategy": "text", 
                "horizontal_strategy": "lines",
                #"snap_tolerance": 3,
                "snap_x_tolerance": 5,
                "snap_y_tolerance": 3,
                "join_x_tolerance": 5,
                "join_y_tolerance": 3,
                "edge_min_length": 3,
                "min_words_vertical": 3,
                "min_words_horizontal": 5,
                "keep_blank_chars": False,
                "text_x_tolerance": 5,
                "text_y_tolerance": 3,
                "intersection_x_tolerance": 5,
                "intersection_y_tolerance": 3,
                }
    extractionSettings["tight"] = {
            "vertical_strategy": "text", 
            "horizontal_strategy": "text",
            "snap_x_tolerance": 1,
            "snap_y_tolerance": 5,
            "join_x_tolerance": 1,
            "join_y_tolerance": 3,
            "edge_min_length": 3,
            "min_words_vertical": 3,
            "min_words_horizontal": 1,
            "keep_blank_chars": False,
            "text_x_tolerance": 1,
            "text_y_tolerance": 3,
            "intersection_x_tolerance": 1,
            "intersection_y_tolerance": 3,
        }


    pdfDFs = {}

    for eSetting in extractionSettings:
        #print(eSetting)
        for page in pdf.pages:
        
            curRow = None
            tableRows = []

            for table in page.extract_tables(extractionSettings[eSetting]):
                for row in table:
                    #print("R", row)
                    
                    if row[0] != None:

                        if curRow != None:
                            curRow = [x.replace("\n", "") if not x is None else "" for x in curRow]
                            tableRows.append(curRow)

                        curRow = row
                        continue

                    if not curRow is None:
                        for xi in range(0, len(curRow)):
                            if row != None and row[xi] != None:

                                if curRow[xi] is None:
                                    curRow[xi] = row[xi]
                                else:
                                    curRow[xi] += "\n" + row[xi]
                    

            pageDF = pd.DataFrame(tableRows)

            if not 0 in pageDF.shape:

                
                if page.page_number in pdfDFs:

                    #print(page.page_number)

                    pageHasSeqCol = hasSeqColumn(pageDF, num_tests=20)
                    existHasSeqCol = hasSeqColumn(pdfDFs[page.page_number], num_tests=20)

                    if pageHasSeqCol and not existHasSeqCol:
                        #print("Replacing PageDF", page.page_number)
                        pass
                    elif pageHasSeqCol and pdfDFs[page.page_number].shape[0] * pdfDFs[page.page_number].shape[1] < pageDF.shape[0]*pageDF.shape[1]:
                        #print("Replacing PageDF", page.page_number)
                        pass
                    else:
                        continue

                pdfDFs[page.page_number] = pageDF



    identified_mirnas = extract_mirnas_sheetsdf(pdfDFs, num_tests=20)

    return pdfDFs, identified_mirnas

def extract_mirnas_from_xml(infile):

    try:
        tree = etree.parse(infile)
    except:
        logging.error("Error loading file XML: {}".format(infile))
        return None, None

    tree = etree.parse(infile)

    #
    ## Checking for Tables
    #
    tables = tree.findall("//table")

    sheetsDF = {}
    for ti, table in enumerate(tables):
        dfs = pd.read_html(etree.tostring(table,method='html'))

        for tii, tdf in enumerate(dfs):
            tname = "Table{}.{}".format(ti, tii)

            fracStrCols = sum([1 if type(x) == str else 0 for x in tdf.columns])/len(tdf.columns)
            
            if fracStrCols > 0 or isinstance(tdf.columns, pd.MultiIndex):
                #print("Removing Columns")
                #print(type(tdf.columns))
                if isinstance(tdf.columns, pd.MultiIndex):
                    #print("Fixing Columns")

                    lvlValues = [""] * len(tdf.columns.get_level_values(0))

                    for x in range(0, len(tdf.columns.levels)):

                        levelValues = tdf.columns.get_level_values(x)

                        if len(set(levelValues)) == 1:
                            continue

                        for vi, v in enumerate(levelValues):

                            if v in lvlValues[vi]:
                                continue

                            if len(lvlValues[vi]) > 0:
                                lvlValues[vi] += " "
                            lvlValues[vi] += v

                    print(lvlValues)

                    tdf.columns = list(tdf.columns.get_level_values(len(tdf.columns.levels)-1))
                    #print(tdf.columns)
                    
                tdf = tdf.T.reset_index().T.reset_index(drop=True)

            sheetsDF[tname] = tdf

            #print(tname)
            #print(tdf)

    print("Sheets", [x for x in sheetsDF])
    identified_mirnas = extract_mirnas_sheetsdf(sheetsDF)

    #
    ## Finding miRNAs in Text!
    #

    pars = tree.findall("//p")
    mirnaTextSearch = re.compile(r'[ATCGU]{15,30}')

    for x in pars:

        if len(x.findall("//table")) > 0:
            continue

        parText = " ".join([y for y in x.itertext()])

        for mirnaMatch in mirnaTextSearch.finditer(parText):
            mirnaSeq = parText[mirnaMatch.span()[0]:mirnaMatch.span()[1]]
            mirnaID = os.path.splitext(os.path.basename(infile))[0] + str(len(identified_mirnas))
            identified_mirnas.add(DocumentMirnaSequence(mirnaID, mirnaSeq, source="PMC_XML_TEXT"))
            

    return sheetsDF, identified_mirnas

def extract_mirnas_sheetsdf(sheetsDF, num_tests=10):

    if len(sheetsDF) == 0:
        return set()

    identified_mirnas = set()

    print("Testing FASTA")
    fasta_mirnas = extract_mirnas_from_pandas_fasta(sheetsDF)
    print("Num FASTA miRNAs", len(fasta_mirnas))
    print("Testing COLUMNS")
    column_mirnas = extract_mirnas_from_pandas_table(sheetsDF, num_tests=num_tests)
    print("Num COLUMNS miRNAs", len(column_mirnas))

    identified_mirnas = fasta_mirnas.union(column_mirnas)

    return identified_mirnas


def extract_mirnas_from_pandas_table(sheetsDF, num_tests=10):
    identified_mirnas = set()
    for sheet in sheetsDF:

        df = sheetsDF[sheet]
        print(sheet, df.shape)
        if 0 in df.shape:
            print("Skipping sheet", sheet)
            continue

        identified_mirnas = identified_mirnas.union(extract_mirnas_from_pandas_tabledf( df, num_tests=num_tests))
    return identified_mirnas

def extract_mirnas_from_pandas_fasta(sheetsDF):
    identified_mirnas = set()
    for sheet in sheetsDF:

        df = sheetsDF[sheet]

        for column in df.columns:

            #print("Testing col", column)

            bestRowSequence = None
            rowSequence = []
            
            for rowIdx, row in df.iterrows():

                fastaStart = fastaStartRE.search(str(row[column]))
                fastaSequenceMir = fastaSequenceMirRE.search(str(row[column]))

                if fastaStart:
                    #print(rowIdx, "fasta start")
                    rowSequence.append("start")
                if fastaSequenceMir:
                    #print(rowIdx, "fasta seq")
                    rowSequence.append("seq")

                if not fastaStart and not fastaSequenceMir:
                    if len(rowSequence) > 0 and (bestRowSequence == None or len(bestRowSequence) > len(rowSequence)):
                        bestRowSequence = rowSequence

                if rowIdx > 20:
                    break
            if len(rowSequence) > 0 and (bestRowSequence == None or len(bestRowSequence) > len(rowSequence)):
                bestRowSequence = rowSequence

            if not bestRowSequence is None:

                count = 0
                bestCount = 0
                for i in range(1, len(bestRowSequence)):

                    if bestRowSequence[i-1] != bestRowSequence[i]:
                        count += 1

                    else:
                        if count > bestCount:
                            bestCount = count

                if count > bestCount:
                            bestCount = count

                if bestCount > 0:
                    print(sheet, column, bestRowSequence)
                    print(sheet, column, bestCount)

                    # extract mirnas

                    startName = None
                    startSeq = None

                    for rowIdx, row in df.iterrows():

                        rowSeq = str(row[column])
                        fastaStart = fastaStartRE.search(rowSeq)
                        fastaSequenceMir = fastaSequenceMirRE.search(rowSeq)

                        if fastaStart:
                            #print(rowIdx, "fasta start")
                            startName = rowSeq[1:]
                        if fastaSequenceMir:
                            #print(rowIdx, "fasta seq")
                            startSeq = rowSeq

                        if not startSeq is None and not startName is None:
                            newMir = DocumentMirnaSequence(startName, startSeq, source="FASTA_COLUMN")
                            identified_mirnas.add(newMir)

                            startName = None
                            startSeq = None
                    
    return identified_mirnas

def merge_mirnas(exmirnas, newmirnas, fileid):

    if newmirnas is None or len(newmirnas) == 0:
        return exmirnas

    fileidPMC = fileid.split("_")[0]
    if fileidPMC.startswith("PMC"):
        fileidPMC = fileidPMC.split(".")[0]
        fileid = fileidPMC

    emirnas = set()
    for x in newmirnas:
        x.add_origin(fileid)
        exmirnas.add(x)
        
    return exmirnas

def merge_sdf(exsdf, newsdf, fileid):

    if newsdf == None or len(newsdf) == 0:
        return exsdf

    exsdf[fileid] = newsdf
    return exsdf

def process_files(infiles):
    found_tables = {}
    identified_mirnas = set()
    for infile in infiles:
        print()
        print()
        print()
        print(infile)

        if "MACOS" in infile:
            logging.warn("Skipping MACOSX file {}".format(infile))
            continue

        sdf = None
        idmirnas = None

        docID = os.path.splitext(os.path.basename(infile))[0]
        
        for foldername in reversed(os.path.dirname(infile).split("/")):
            if foldername.startswith("PMC"):
                docID = foldername.split("_")[0]
                break

        if infile.upper().endswith(".XLSX"):
            sdf, idmirnas = extract_mirnas_from_excel(infile)
            print("XLSX", len(idmirnas))

            if not idmirnas is None:
                identified_mirnas = identified_mirnas.union(idmirnas)

        elif infile.upper().endswith(".XML"):
            sdf, idmirnas = extract_mirnas_from_xml(infile)

            if not idmirnas is None:
                identified_mirnas = identified_mirnas.union(idmirnas)

        elif infile.upper().endswith(".PDF"):
            sdf, idmirnas = extract_mirnas_from_pdf(infile)
            if not idmirnas is None:
                identified_mirnas = identified_mirnas.union(idmirnas)

        elif infile.upper().endswith(".DOCX"):
            sdf, idmirnas = extract_mirnas_from_document(infile)

            if not idmirnas is None:
                identified_mirnas = identified_mirnas.union(idmirnas)

        elif infile.upper().endswith(".ZIP"):

            continue

            outfolder = os.path.dirname(infile) + "/" + os.path.splitext(os.path.basename(infile))[0]
            print("Outfolder", outfolder)

            try:
                with zipfile.ZipFile(infile) as z:
                    z.extractall(outfolder)
            except:
                print("Invalid file")

            sdf, idmirnas = process_files(find_relevant_files(outfolder))
            
        
        identified_mirnas = merge_mirnas(identified_mirnas, idmirnas, docID)
        found_tables = merge_sdf(found_tables, sdf, docID )


    return sdf, set(list(identified_mirnas))
    

